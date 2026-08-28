"""P2: 需要 fixture 的集成测试"""

import os
import sys
import json
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest
from PIL import Image, ImageDraw


# ===== Fixture helpers =====

@pytest.fixture
def sample_docx(tmp_path):
    """生成一个带图片的 .docx 文件，放在板块目录下"""
    from docx import Document
    from docx.shared import Inches

    img_path = tmp_path / "_img.png"
    img = Image.new("RGB", (100, 100), (255, 0, 0))
    img.save(str(img_path))

    doc = Document()
    doc.add_paragraph("标题文本")
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(1))
    doc.add_paragraph("正文文本")

    section_dir = tmp_path / "inputs"
    section_dir.mkdir(parents=True)
    path = section_dir / "sample.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def text_only_docx(tmp_path):
    """生成纯文本 .docx"""
    from docx import Document
    doc = Document()
    doc.add_paragraph("纯文本内容")
    section_dir = tmp_path / "inputs"
    section_dir.mkdir(parents=True)
    path = section_dir / "text_only.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def test_images(tmp_path):
    """生成两张测试图片"""
    img1 = Image.new("RGB", (50, 50), (100, 100, 100))
    p1 = tmp_path / "gold.png"
    img1.save(str(p1))

    img2 = Image.new("RGB", (50, 50), (100, 100, 100))
    p2 = tmp_path / "actual.png"
    img2.save(str(p2))

    img3 = Image.new("RGB", (50, 50), (200, 100, 100))
    p3 = tmp_path / "diff_actual.png"
    img3.save(str(p3))

    img4 = Image.new("RGB", (60, 50), (100, 100, 100))
    p4 = tmp_path / "size_mismatch.png"
    img4.save(str(p4))

    return {
        "identical_gold": str(p1),
        "identical_actual": str(p2),
        "different_actual": str(p3),
        "size_mismatch": str(p4),
    }


@pytest.fixture
def template_setup(tmp_path):
    """生成模板配置目录 + 相关路径"""
    for tid in ["1_本周头条", "4_一句话"]:
        d = tmp_path / "templates" / tid
        d.mkdir(parents=True)
        (d / "config.json").write_text(json.dumps({
            "layout_mode": "templateA",
        }))
    return {
        "templates_root": str(tmp_path / "templates"),
        "inputs_root": str(tmp_path / "inputs"),
        "workspace_root": str(tmp_path / "outputs"),
    }


# ===== Tests =====

class TestExportDocxImages:
    def test_export_images(self, sample_docx, tmp_path):
        from docx_list_to_json import export_docx_images
        result = export_docx_images(sample_docx, str(tmp_path / "images"))
        assert len(result) > 0
        for rel, abspath in result.items():
            assert os.path.exists(abspath)

    def test_no_images(self, text_only_docx, tmp_path):
        from docx_list_to_json import export_docx_images
        result = export_docx_images(text_only_docx, str(tmp_path / "images"))
        assert result == {}

    def test_with_prefix(self, sample_docx, tmp_path):
        from docx_list_to_json import export_docx_images
        result = export_docx_images(sample_docx, str(tmp_path / "images"), name_prefix="test")
        for rel, abspath in result.items():
            filename = os.path.basename(abspath)
            assert filename.startswith("test_")


class TestDocxListToJson:
    def test_basic(self, sample_docx, template_setup, tmp_path):
        from docx_list_to_json import docx_list_to_json, discover_template_config_map

        template_map = discover_template_config_map(template_setup["templates_root"])
        output_dir = template_setup["workspace_root"]
        os.makedirs(output_dir, exist_ok=True)

        elements = docx_list_to_json(
            source_items=[{"kind": "word", "path": sample_docx, "section_name": "1_本周头条"}],
            template_config_map=template_map,
            input_root_dir=template_setup["inputs_root"],
            doc_workspace_root_dir=output_dir,
            image_root_dir=str(tmp_path / "_images"),
        )
        assert len(elements) > 0
        assert all("doc_name" in e for e in elements)
        assert all("template_id" in e for e in elements)

    def test_image_mode(self, template_setup, tmp_path):
        from docx_list_to_json import docx_list_to_json, discover_template_config_map, discover_input_sources

        # 创建图片模式输入
        section_dir = Path(template_setup["inputs_root"]) / "4_一句话"
        section_dir.mkdir(parents=True)
        img = Image.new("RGB", (50, 50), (0, 255, 0))
        img.save(str(section_dir / "test.png"))

        template_map = discover_template_config_map(template_setup["templates_root"])
        output_dir = template_setup["workspace_root"]
        os.makedirs(output_dir, exist_ok=True)

        sources = discover_input_sources(template_setup["inputs_root"])

        elements = docx_list_to_json(
            source_items=sources,
            template_config_map=template_map,
            input_root_dir=template_setup["inputs_root"],
            doc_workspace_root_dir=output_dir,
            image_root_dir=str(tmp_path / "_images"),
        )
        assert len(elements) >= 2  # image + text pair

    def test_multiple_docs(self, sample_docx, template_setup, tmp_path):
        from docx_list_to_json import docx_list_to_json, discover_template_config_map
        from shutil import copy2

        doc2_path = Path(template_setup["inputs_root"]) / "1_本周头条" / "doc2.docx"
        doc2_path.parent.mkdir(parents=True, exist_ok=True)
        copy2(sample_docx, str(doc2_path))

        template_map = discover_template_config_map(template_setup["templates_root"])
        output_dir = template_setup["workspace_root"]
        os.makedirs(output_dir, exist_ok=True)

        elements = docx_list_to_json(
            source_items=[
                {"kind": "word", "path": sample_docx, "section_name": "1_本周头条"},
                {"kind": "word", "path": str(doc2_path), "section_name": "1_本周头条"},
            ],
            template_config_map=template_map,
            input_root_dir=template_setup["inputs_root"],
            doc_workspace_root_dir=output_dir,
            image_root_dir=str(tmp_path / "_images"),
        )
        assert len(elements) > 0

        # 每个文档应生成独立的任务缓存 JSON
        output_jsons = list((Path(output_dir) / "caches").glob("*.json"))
        assert len(output_jsons) >= 2


class TestComparePageImage:
    def test_identical(self, test_images, tmp_path):
        from compare_snapshot import compare_page_image
        diff_out = tmp_path / "diff.png"
        result = compare_page_image(
            Path(test_images["identical_actual"]),
            Path(test_images["identical_gold"]),
            diff_out,
            threshold=10,
        )
        assert result.get("diff_ratio", 1) == 0

    def test_different(self, test_images, tmp_path):
        from compare_snapshot import compare_page_image
        diff_out = tmp_path / "diff.png"
        result = compare_page_image(
            Path(test_images["different_actual"]),
            Path(test_images["identical_gold"]),
            diff_out,
            threshold=10,
        )
        assert result.get("diff_ratio", 0) > 0

    def test_size_mismatch(self, test_images, tmp_path):
        from compare_snapshot import compare_page_image
        diff_out = tmp_path / "diff.png"
        result = compare_page_image(
            Path(test_images["size_mismatch"]),
            Path(test_images["identical_gold"]),
            diff_out,
            threshold=10,
        )
        assert result.get("size_mismatch") is True

    def test_not_exists(self, tmp_path):
        from compare_snapshot import compare_page_image
        diff_out = tmp_path / "diff.png"
        result = compare_page_image(
            Path(tmp_path / "nonexistent.png"),
            Path(tmp_path / "nonexistent2.png"),
            diff_out,
        )
        assert "error" in result
